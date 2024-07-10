from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import redirect, render
from django.urls import reverse_lazy
from django.views.generic import DetailView, ListView, View
from django.views.generic.edit import CreateView, DeleteView, UpdateView

from apps.students.models import Student

from .forms import InvoiceItemFormset, InvoiceReceiptFormSet, Invoices
from .models import Invoice, InvoiceItem, Receipt


class InvoiceListView(LoginRequiredMixin, ListView):
    model = Invoice


class InvoiceCreateView(LoginRequiredMixin, CreateView):
    model = Invoice
    fields = "__all__"
    success_url = "/finance/list"

    def get_context_data(self, **kwargs):
        context = super(InvoiceCreateView, self).get_context_data(**kwargs)
        if self.request.POST:
            context["items"] = InvoiceItemFormset(
                self.request.POST, prefix="invoiceitem_set"
            )
        else:
            context["items"] = InvoiceItemFormset(prefix="invoiceitem_set")
        return context

    def form_valid(self, form):
        context = self.get_context_data()
        formset = context["items"]
        self.object = form.save()
        if self.object.id != None:
            if form.is_valid() and formset.is_valid():
                formset.instance = self.object
                formset.save()
        return super().form_valid(form)


class InvoiceDetailView(LoginRequiredMixin, DetailView):
    model = Invoice
    fields = "__all__"

    def get_context_data(self, **kwargs):
        context = super(InvoiceDetailView, self).get_context_data(**kwargs)
        context["receipts"] = Receipt.objects.filter(invoice=self.object)
        context["items"] = InvoiceItem.objects.filter(invoice=self.object)
        return context


class InvoiceUpdateView(LoginRequiredMixin, UpdateView):
    model = Invoice
    fields = ["student", "session", "term", "class_for", "balance_from_previous_term"]

    def get_context_data(self, **kwargs):
        context = super(InvoiceUpdateView, self).get_context_data(**kwargs)
        if self.request.POST:
            context["receipts"] = InvoiceReceiptFormSet(
                self.request.POST, instance=self.object
            )
            context["items"] = InvoiceItemFormset(
                self.request.POST, instance=self.object
            )
        else:
            context["receipts"] = InvoiceReceiptFormSet(instance=self.object)
            context["items"] = InvoiceItemFormset(instance=self.object)
        return context

    def form_valid(self, form):
        context = self.get_context_data()
        formset = context["receipts"]
        itemsformset = context["items"]
        if form.is_valid() and formset.is_valid() and itemsformset.is_valid():
            form.save()
            formset.save()
            itemsformset.save()
        return super().form_valid(form)


class InvoiceDeleteView(LoginRequiredMixin, DeleteView):
    model = Invoice
    success_url = reverse_lazy("invoice-list")


class ReceiptCreateView(LoginRequiredMixin, CreateView):
    model = Receipt
    fields = ["amount_paid", "date_paid", "comment"]
    success_url = reverse_lazy("invoice-list")

    def form_valid(self, form):
        obj = form.save(commit=False)
        invoice = Invoice.objects.get(pk=self.request.GET["invoice"])
        obj.invoice = invoice
        obj.save()
        return redirect("invoice-list")

    def get_context_data(self, **kwargs):
        context = super(ReceiptCreateView, self).get_context_data(**kwargs)
        invoice = Invoice.objects.get(pk=self.request.GET["invoice"])
        context["invoice"] = invoice
        return context


class ReceiptUpdateView(LoginRequiredMixin, UpdateView):
    model = Receipt
    fields = ["amount_paid", "date_paid", "comment"]
    success_url = reverse_lazy("invoice-list")


class ReceiptDeleteView(LoginRequiredMixin, DeleteView):
    model = Receipt
    success_url = reverse_lazy("invoice-list")


@login_required
def bulk_invoice(request):
    return render(request, "finance/bulk_invoice.html")




import csv


from django.http import HttpResponse


# class DownloadCSVViewdownloadcsv(LoginRequiredMixin, View):
#     def get(self, request, *args, **kwargs):
#         response = HttpResponse(content_type="text/csv")
#         response["Content-Disposition"] = 'attachment; filename="student_template.csv"'

#         writer = csv.writer(response)
#         writer.writerow(
#             [
#                 "registration_number",
#                 "surname",
#                 "firstname1",
#                 "other_names1",
#                 "gender1",
#                 "parent_number1",
#                 "address1",
#                 "current_class",
#             ]
#         )

#         return response



########### download the word format:: good but some content not visible

# from django.shortcuts import get_object_or_404, HttpResponse
# from django.template.loader import render_to_string
# from django.http import HttpResponse
# from .models import Invoice, InvoiceItem, Receipt  # Adjust based on your actual models
# from docx import Document
# from docx.shared import Pt
# from django.utils import timezone
# import os

# def download_invoice(request, invoice_id):
#     # Fetch invoice details
#     invoice = get_object_or_404(Invoice, pk=invoice_id)
#     items = InvoiceItem.objects.filter(invoice=invoice)
#     receipts = Receipt.objects.filter(invoice=invoice)

#     # Create a new Document
#     doc = Document()
#     doc.add_heading(f'Invoice #{invoice.id}', 0)

#     # Add invoice details
#     doc.add_paragraph(f'Session: {invoice.session}')
#     doc.add_paragraph(f'Term: {invoice.term}')
#     doc.add_paragraph(f'Class: {invoice.class_for}')
#     doc.add_paragraph(f'Status: {invoice.get_status_display()}')

#     doc.add_paragraph('Expected Balance: {}'.format(invoice.balance))

#     # Add table for Invoice Breakdown
#     doc.add_heading('Invoice Breakdown', level=1)
#     table = doc.add_table(rows=1, cols=3)
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'S/N'
#     hdr_cells[1].text = 'Description'
#     hdr_cells[2].text = 'Amount'

#     for i, item in enumerate(items, start=1):
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(i)
#         row_cells[1].text = item.description
#         row_cells[2].text = str(item.amount)

#     # Add totals
#     doc.add_paragraph(f'Total Amount this term: {invoice.amount_payable}')
#     doc.add_paragraph(f'Balance from previous term: {invoice.balance_from_previous_term}')
#     doc.add_paragraph(f'Total Amount Payable: {invoice.total_amount_payable}')
#     doc.add_paragraph(f'Total Amount Paid: {invoice.total_amount_paid}')

#     # Add Payment History
#     doc.add_heading('Payment History', level=1)
#     table = doc.add_table(rows=1, cols=4)
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'S/N'
#     hdr_cells[1].text = 'Amount Paid'
#     hdr_cells[2].text = 'Date Paid'
#     hdr_cells[3].text = 'Comment'

#     for i, receipt in enumerate(receipts, start=1):
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(i)
#         row_cells[1].text = str(receipt.amount_paid)
#         row_cells[2].text = receipt.date_paid.strftime('%Y-%m-%d')
#         row_cells[3].text = receipt.comment

#     # Save the document to a BytesIO object
#     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = f'attachment; filename=invoice_{invoice.id}.docx'
#     doc.save(response)
#     return response



##### data is visisble in table format and same some data is missing

# from django.shortcuts import get_object_or_404, HttpResponse
# from django.http import HttpResponse
# from .models import Invoice, InvoiceItem, Receipt  # Adjust based on your actual models
# from docx import Document
# from docx.shared import Pt
# from django.utils import timezone

# def download_invoice(request, invoice_id):
#     # Fetch invoice details
#     invoice = get_object_or_404(Invoice, pk=invoice_id)
#     items = InvoiceItem.objects.filter(invoice=invoice)
#     receipts = Receipt.objects.filter(invoice=invoice)

#     # Create a new Document
#     doc = Document()
#     doc.add_heading(f'Invoice #{invoice.id}', 0)

#     # Add invoice details in a 2-column table
#     table = doc.add_table(rows=5, cols=2)
#     table.style = 'Table Grid'

#     row = table.rows[0]
#     row.cells[0].text = 'name'
#     row.cells[1].text = str(invoice)

#     row = table.rows[1]
#     row.cells[0].text = 'Session'
#     row.cells[1].text = str(invoice.session)  # Ensure session is converted to string

#     row = table.rows[2]
#     row.cells[0].text = 'Term'
#     row.cells[1].text = str(invoice.term)  # Ensure term is converted to string

#     row = table.rows[3]
#     row.cells[0].text = 'Class'
#     row.cells[1].text = str(invoice.class_for)  # Ensure class_for is converted to string

#     row = table.rows[4]
#     row.cells[0].text = 'Status'
#     row.cells[1].text = str(invoice.get_status_display())  # Ensure status is converted to string

#     doc.add_paragraph()

#     # Add Expected Balance
#     doc.add_paragraph(f'Expected Balance: {invoice.balance}')

#     doc.add_paragraph()

#     # Add Invoice Breakdown table
#     doc.add_heading('Invoice Breakdown', level=1)
#     table = doc.add_table(rows=1, cols=3)
#     table.style = 'Table Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'S/N'
#     hdr_cells[1].text = 'Description'
#     hdr_cells[2].text = 'Amount'

#     for i, item in enumerate(items, start=1):
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(i)
#         row_cells[1].text = item.description
#         row_cells[2].text = str(item.amount)

#     # Add totals
#     doc.add_paragraph(f'Total Amount this term: {invoice.amount_payable}')
#     doc.add_paragraph(f'Balance from previous term: {invoice.balance_from_previous_term}')
#     doc.add_paragraph(f'Total Amount Payable: {invoice.total_amount_payable}')
#     doc.add_paragraph(f'Total Amount Paid: {invoice.total_amount_paid}')

#     doc.add_paragraph()

#     # Add Payment History
#     doc.add_heading('Payment History', level=1)
#     table = doc.add_table(rows=1, cols=4)
#     table.style = 'Table Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'S/N'
#     hdr_cells[1].text = 'Amount Paid'
#     hdr_cells[2].text = 'Date Paid'
#     hdr_cells[3].text = 'Comment'

#     for i, receipt in enumerate(receipts, start=1):
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(i)
#         row_cells[1].text = str(receipt.amount_paid)
#         row_cells[2].text = receipt.date_paid.strftime('%Y-%m-%d')
#         row_cells[3].text = receipt.comment

#     # Save the document to a response
#     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = f'attachment; filename=invoice_{invoice.id}.docx'
#     doc.save(response)
#     return response



########### download invoice successfully word format working good condition #############

# from django.shortcuts import get_object_or_404, HttpResponse
# from .models import Invoice, InvoiceItem, Receipt  # Adjust based on your actual models
# from docx import Document

# def download_invoice(request, invoice_id):
#     # Fetch invoice details
#     invoice = get_object_or_404(Invoice, pk=invoice_id)
#     items = InvoiceItem.objects.filter(invoice=invoice)
#     receipts = Receipt.objects.filter(invoice=invoice)

#     # Create a new Document
#     doc = Document()
#     doc.add_heading(f'Invoice #{invoice.id}', 0)

#     # Add invoice details in a 2-column table
#     table = doc.add_table(rows=5, cols=2)
#     table.style = 'Table Grid'

#     row = table.rows[0]
#     row.cells[0].text = 'Invoice'
#     row.cells[1].text = str(invoice)

#     row = table.rows[1]
#     row.cells[0].text = 'Session'
#     row.cells[1].text = str(invoice.session)  # Ensure session is converted to string

#     row = table.rows[2]
#     row.cells[0].text = 'Term'
#     row.cells[1].text = str(invoice.term)  # Ensure term is converted to string

#     row = table.rows[3]
#     row.cells[0].text = 'Class'
#     row.cells[1].text = str(invoice.class_for)  # Ensure class_for is converted to string

#     row = table.rows[4]
#     row.cells[0].text = 'Status'
#     row.cells[1].text = str(invoice.get_status_display())  # Ensure status is converted to string

#     doc.add_paragraph()

#     # Add Expected Balance
#     doc.add_paragraph(f'Expected Balance: {invoice.balance()}')  # Call the balance method

#     doc.add_paragraph()

#     # Add Invoice Breakdown table
#     doc.add_heading('Invoice Breakdown', level=1)
#     table = doc.add_table(rows=1, cols=3)
#     table.style = 'Table Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'S/N'
#     hdr_cells[1].text = 'Description'
#     hdr_cells[2].text = 'Amount'

#     for i, item in enumerate(items, start=1):
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(i)
#         row_cells[1].text = item.description
#         row_cells[2].text = str(item.amount)

#     # Add totals
#     doc.add_paragraph(f'Total Amount this term: {invoice.amount_payable()}')  # Call the amount_payable method
#     doc.add_paragraph(f'Balance from previous term: {invoice.balance_from_previous_term}')  # Already a string
#     doc.add_paragraph(f'Total Amount Payable: {invoice.total_amount_payable()}')  # Call the total_amount_payable method
#     doc.add_paragraph(f'Total Amount Paid: {invoice.total_amount_paid()}')  # Call the total_amount_paid method

#     doc.add_paragraph()

#     # Add Payment History
#     doc.add_heading('Payment History', level=1)
#     table = doc.add_table(rows=1, cols=4)
#     table.style = 'Table Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'S/N'
#     hdr_cells[1].text = 'Amount Paid'
#     hdr_cells[2].text = 'Date Paid'
#     hdr_cells[3].text = 'Comment'

#     for i, receipt in enumerate(receipts, start=1):
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(i)
#         row_cells[1].text = str(receipt.amount_paid)
#         row_cells[2].text = receipt.date_paid.strftime('%Y-%m-%d')
#         row_cells[3].text = receipt.comment

#     # Save the document to a response
#     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = f'attachment; filename=invoice_{invoice.id}.docx'
#     doc.save(response)
#     return response






from django.shortcuts import get_object_or_404, HttpResponse
from .models import Invoice, InvoiceItem, Receipt  # Adjust based on your actual models
from docx import Document

def download_invoice(request, invoice_id):
    # Fetch invoice details
    invoice = get_object_or_404(Invoice, pk=invoice_id)
    items = InvoiceItem.objects.filter(invoice=invoice)
    receipts = Receipt.objects.filter(invoice=invoice)

    # Create a new Document
    doc = Document()
    doc.add_heading(f'Invoice #{invoice.id}', level=1)

    # Add invoice details in a 2-column table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    row = table.rows[0]
    row.cells[0].text = 'Invoice'
    row.cells[1].text = str(invoice)

    row = table.rows[1]
    row.cells[0].text = 'Session'
    row.cells[1].text = str(invoice.session)

    row = table.rows[2]
    row.cells[0].text = 'Term'
    row.cells[1].text = str(invoice.term)

    row = table.rows[3]
    row.cells[0].text = 'Class'
    row.cells[1].text = str(invoice.class_for)

    row = table.rows[4]
    row.cells[0].text = 'Status'
    row.cells[1].text = str(invoice.get_status_display())

    doc.add_paragraph()

    # Add Expected Balance
    doc.add_paragraph(f'Expected Balance: {invoice.balance()}')

    doc.add_paragraph()

    # Add Invoice Breakdown table
    doc.add_heading('Invoice Breakdown', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S/N'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Amount'

    for i, item in enumerate(items, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = item.description
        row_cells[2].text = str(item.amount)

    # Add totals
    doc.add_paragraph(f'Total Amount this term: {invoice.amount_payable()}')
    doc.add_paragraph(f'Balance from previous term: {invoice.balance_from_previous_term}')
    doc.add_paragraph(f'Total Amount Payable: {invoice.total_amount_payable()}')
    doc.add_paragraph(f'Total Amount Paid: {invoice.total_amount_paid()}')

    doc.add_paragraph()

    # Add Payment History
    doc.add_heading('Payment History', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S/N'
    hdr_cells[1].text = 'Amount Paid'
    hdr_cells[2].text = 'Date Paid'
    hdr_cells[3].text = 'Comment'

    for i, receipt in enumerate(receipts, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(receipt.amount_paid)
        row_cells[2].text = receipt.date_paid.strftime('%Y-%m-%d')  # Format date as needed
        row_cells[3].text = receipt.comment

    # Save the document to a response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=invoice_{invoice.id}.docx'
    doc.save(response)
    return response












































### ### ###

# from django.shortcuts import render, redirect
# from .forms import StudentForm

# def student_create(request):
#     if request.method == 'POST':
#         form = StudentForm(request.POST, request.FILES)
#         if form.is_valid():
#             form.save()
#             return redirect('student_list')
#     else:
#         form = StudentForm()
#     return render(request, 'student_form.html', {'form': form})


from django.shortcuts import render, get_object_or_404
from .models import Student

def student_detail(request, student_id):
    student = get_object_or_404(Student, id=student_id)
    return render(request, 'student_detail.html', {'object': student})

